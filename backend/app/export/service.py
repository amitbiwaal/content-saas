"""Schema & Export service (PRD §9.10, §10 export, FR-10.1..10.4).

Turns a draft (the ``{"sections": [...], "word_count": int}`` payload produced by
the Article Writer, PRD §7/§14) plus its project brief into portable export
formats:

* :func:`build_jsonld` — JSON-LD appropriate to the content type (FR-10.1):
  ``Article`` by default, ``Review`` for review content, and a ``FAQPage`` graph
  for FAQ content. The ``FAQPage`` block is explicitly flagged parse-value-only
  (FR-10.4): FAQ schema is emitted for AI/answer-engine parsing, never as a SERP
  rich-result claim (FAQPage rich results were deprecated, PRD §11).
* :func:`to_markdown` — portable Markdown (FR-10.2).
* :func:`to_html` — HTML with inline house-style typography (PRD §11 / §15): the
  body is locked to ``16px`` / ``#334155`` inline; only headers, byline and table
  cells may differ (FR-10.2, HTML carries inline house-style typography).
* :func:`to_docx` — DOCX bytes via ``python-docx`` if installed, else a clear
  :class:`NotImplementedError` (the dependency is imported lazily and the module
  degrades rather than failing to import, FR-10.2).
* :func:`wordpress_publish` — a STUB returning a would-publish payload; the real
  WordPress REST integration is wired manually later (FR-10.2, PRD §15).

These functions are PURE (draft/project dicts -> str/dict/bytes) so they are
unit-testable without a DB or any provider keys. No LLM calls are made here:
export is deterministic serialisation, so the zero-API-key contract holds
trivially. DB persistence (latest Score/Draft lookups, the publish gate) lives in
the router (:mod:`app.export.router`).
"""

from __future__ import annotations

import html as _html
import re
from typing import Any

# Body typography locked by the house rules (PRD §11 / §15): 16px / #334155.
# Only headers, byline and table cells may differ.
BODY_FONT_SIZE = "16px"
BODY_COLOR = "#334155"
_BODY_STYLE = f"font-size:{BODY_FONT_SIZE};color:{BODY_COLOR};line-height:1.7;"

# Content-type aliases -> canonical schema.org type (FR-10.1).
_REVIEW_TYPES = {"review", "product_review", "productreview", "comparison"}
_FAQ_TYPES = {"faq", "faqpage", "faq_page", "qa", "q&a"}


# --------------------------------------------------------------------------- #
# Draft normalisation (accept the Article Writer payload, an ORM-row dict, or a
# bare list of sections — siblings produce slightly different shapes).
# --------------------------------------------------------------------------- #
def _sections(draft: Any) -> list[dict]:
    """Return the ordered section dicts from a draft payload, defensively.

    Each section is normalised to ``{"heading", "level", "markdown"}``; unknown
    shapes are coerced rather than raising, so a partial draft still exports.
    """
    if isinstance(draft, dict):
        raw = draft.get("sections")
    elif isinstance(draft, list):
        raw = draft
    else:
        raw = None
    if not isinstance(raw, list):
        return []

    out: list[dict] = []
    for item in raw:
        if not isinstance(item, dict):
            continue
        heading = ""
        for key in ("heading", "title", "text", "name"):
            value = item.get(key)
            if isinstance(value, str) and value.strip():
                heading = value.strip()
                break
        level_raw = item.get("level", 2)
        level = level_raw if isinstance(level_raw, int) and not isinstance(level_raw, bool) else 2
        if level < 2:
            level = 2
        if level > 3:
            level = 3
        body = item.get("markdown")
        if not isinstance(body, str):
            body = item.get("body") if isinstance(item.get("body"), str) else ""
        out.append({"heading": heading, "level": level, "markdown": body.strip()})
    return out


def _title(draft: Any, project: dict | None) -> str:
    """Resolve the document/H1 title: explicit draft title, else the brief topic."""
    if isinstance(draft, dict):
        for key in ("title", "h1", "headline"):
            value = draft.get(key)
            if isinstance(value, str) and value.strip():
                return value.strip()
    if isinstance(project, dict):
        for key in ("topic", "keyword"):
            value = project.get(key)
            if isinstance(value, str) and value.strip():
                return value.strip()
    return "Untitled"


def _plain_text(markdown: str) -> str:
    """Strip the lightweight Markdown we emit to plain text (for JSON-LD body)."""
    text = re.sub(r"`{1,3}", "", markdown or "")
    text = re.sub(r"[*_#>]+", "", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", text)  # links -> anchor text
    return re.sub(r"[ \t]+", " ", text).strip()


def _faq_pairs(sections: list[dict]) -> list[tuple[str, str]]:
    """Pull (question, answer) pairs out of FAQ-ish sections.

    A section heading ending in '?' is treated as a question with its body as the
    answer; otherwise each non-empty markdown line ending in '?' starts a Q with
    the following lines as the A. Best-effort so FAQ JSON-LD is still produced.
    """
    pairs: list[tuple[str, str]] = []
    for section in sections:
        heading = section["heading"]
        body = section["markdown"]
        if heading.endswith("?") and body:
            pairs.append((heading, _plain_text(body)))
            continue
        # Fall back to scanning the body for "Q?\nA" runs.
        question: str | None = None
        answer: list[str] = []
        for line in body.splitlines():
            stripped = line.strip().lstrip("#*->").strip()
            if not stripped:
                continue
            if stripped.endswith("?"):
                if question and answer:
                    pairs.append((question, _plain_text(" ".join(answer))))
                question, answer = stripped, []
            elif question:
                answer.append(stripped)
        if question and answer:
            pairs.append((question, _plain_text(" ".join(answer))))
    return pairs


# --------------------------------------------------------------------------- #
# FR-10.1 / FR-10.4 — JSON-LD
# --------------------------------------------------------------------------- #
def build_jsonld(content_type: str, draft: dict, project: dict) -> dict:
    """Build JSON-LD appropriate to the content type (FR-10.1).

    Args:
        content_type: Free-form content type (e.g. ``"article"``, ``"review"``,
            ``"faq"``). Mapped to a schema.org ``@type``: ``Review`` for review
            content, ``FAQPage`` for FAQ content, else ``Article``.
        draft: Draft payload (Article Writer shape or an ORM-row dict).
        project: Project brief dict (``topic``/``keyword``/``website``/...).

    Returns:
        A JSON-LD ``dict``. For FAQ content the document is a ``FAQPage`` and is
        flagged parse-value-only (FR-10.4): ``"isPartOf": {"@type": "WebPage"}``
        plus a ``x-contentos-parse-value-only: true`` marker so downstream code
        never claims a FAQPage SERP rich result (deprecated, PRD §11).
    """
    kind = (content_type or "").strip().lower()
    sections = _sections(draft)
    title = _title(draft, project)
    website = (project or {}).get("website") or ""
    keyword = (project or {}).get("keyword") or ""
    body_text = "\n\n".join(s["markdown"] for s in sections if s["markdown"])
    plain_body = _plain_text(body_text)

    if kind in _FAQ_TYPES:
        pairs = _faq_pairs(sections)
        return {
            "@context": "https://schema.org",
            "@type": "FAQPage",
            "name": title,
            "about": keyword or title,
            # FR-10.4: FAQ schema is for AI/answer-engine parsing only, NOT a
            # SERP rich-result claim (FAQPage rich results deprecated, PRD §11).
            "x-contentos-parse-value-only": True,
            "isPartOf": {"@type": "WebPage", "url": website} if website else {"@type": "WebPage"},
            "mainEntity": [
                {
                    "@type": "Question",
                    "name": q,
                    "acceptedAnswer": {"@type": "Answer", "text": a},
                }
                for q, a in pairs
            ],
        }

    if kind in _REVIEW_TYPES:
        return {
            "@context": "https://schema.org",
            "@type": "Review",
            "name": title,
            "headline": title,
            "about": keyword or title,
            "itemReviewed": {"@type": "Thing", "name": keyword or title},
            "reviewBody": plain_body,
            "author": {"@type": "Organization", "name": website or "Editorial Team"},
            "publisher": {"@type": "Organization", "name": website or "ContentOS"},
        }

    # Default: Article (FR-10.1).
    return {
        "@context": "https://schema.org",
        "@type": "Article",
        "headline": title,
        "name": title,
        "about": keyword or title,
        "articleBody": plain_body,
        "author": {"@type": "Organization", "name": website or "Editorial Team"},
        "publisher": {"@type": "Organization", "name": website or "ContentOS"},
    }


# --------------------------------------------------------------------------- #
# FR-10.2 — Markdown
# --------------------------------------------------------------------------- #
def to_markdown(draft: dict) -> str:
    """Render the draft as portable Markdown (FR-10.2).

    The document H1 (the draft title) is emitted first, then one heading +
    body block per section at its outline level (H2/H3). Returns a non-empty
    string for any draft that has a title or at least one section.
    """
    sections = _sections(draft)
    title = _title(draft, None)

    lines: list[str] = [f"# {title}", ""]
    for section in sections:
        marker = "#" * section["level"]
        lines.append(f"{marker} {section['heading']}".rstrip())
        lines.append("")
        if section["markdown"]:
            lines.append(section["markdown"])
            lines.append("")
    return "\n".join(lines).rstrip() + "\n"


# --------------------------------------------------------------------------- #
# FR-10.2 / PRD §11 — HTML with inline house-style typography
# --------------------------------------------------------------------------- #
def _safe_href(url: str) -> str:
    """Allow only http/https/mailto/relative link targets (block javascript:/data:).

    Draft markdown is model-generated; without this, a link like
    ``[x](javascript:...)`` would emit a clickable script URL in the HTML /
    Gutenberg export (stored XSS).
    """
    stripped = url.strip()
    if re.match(r"^(https?:|mailto:|/|#)", stripped, re.IGNORECASE):
        return _html.escape(stripped, quote=True)
    return "#"


def _md_inline_to_html(text: str) -> str:
    """Escape and apply minimal inline Markdown (bold/italic/links) for HTML."""
    out = _html.escape(text)
    out = re.sub(
        r"\[([^\]]+)\]\(([^)]+)\)",
        lambda m: f'<a href="{_safe_href(m.group(2))}">{m.group(1)}</a>',
        out,
    )
    out = re.sub(r"\*\*([^*]+)\*\*", r"<strong>\1</strong>", out)
    out = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"<em>\1</em>", out)
    return out


def _markdown_body_to_html(markdown: str) -> str:
    """Render a section's Markdown body to HTML paragraphs with body typography.

    Blank-line-separated blocks become ``<p>`` elements; each paragraph carries
    the locked body style inline (16px / #334155, PRD §11).
    """
    blocks = [b.strip() for b in re.split(r"\n\s*\n", markdown or "") if b.strip()]
    parts: list[str] = []
    for block in blocks:
        inner = _md_inline_to_html(block.replace("\n", " "))
        parts.append(f'<p style="{_BODY_STYLE}">{inner}</p>')
    return "\n".join(parts)


def to_html(draft: dict, *, inline_typography: bool = True) -> str:
    """Render the draft as an HTML fragment (FR-10.2, PRD §11 / §15).

    Args:
        draft: Draft payload (Article Writer shape or an ORM-row dict).
        inline_typography: When True (default), the body paragraphs carry the
            locked house style inline (``font-size:16px;color:#334155``). Headers,
            byline and table cells are left to differ per the house rule (PRD §11).
            Set False to emit semantic HTML without the inline body style.

    Returns:
        A non-empty HTML string wrapped in an ``<article>`` element.
    """
    sections = _sections(draft)
    title = _title(draft, None)

    body_attr = f' style="{_BODY_STYLE}"' if inline_typography else ""

    parts: list[str] = [f"<article{body_attr}>", f"<h1>{_html.escape(title)}</h1>"]
    for section in sections:
        tag = f"h{section['level']}"
        # Headers may differ from body typography (PRD §11) — no inline body style.
        parts.append(f"<{tag}>{_html.escape(section['heading'])}</{tag}>")
        if section["markdown"]:
            if inline_typography:
                parts.append(_markdown_body_to_html(section["markdown"]))
            else:
                blocks = [
                    b.strip() for b in re.split(r"\n\s*\n", section["markdown"]) if b.strip()
                ]
                for block in blocks:
                    parts.append(f"<p>{_md_inline_to_html(block.replace(chr(10), ' '))}</p>")
    parts.append("</article>")
    return "\n".join(parts)


# --------------------------------------------------------------------------- #
# PRD §15 — WordPress: Gutenberg blocks
# --------------------------------------------------------------------------- #
def to_gutenberg(draft: dict) -> str:
    """Render the draft body as WordPress **Gutenberg block** markup (PRD §15).

    Each outline heading becomes a ``wp:heading`` block and each Markdown
    paragraph a ``wp:paragraph`` block — the block-comment format WordPress's
    block editor expects. The H1/title is the *post* title (set separately), so
    it is intentionally not emitted here.
    """
    parts: list[str] = []
    for section in _sections(draft):
        level = section["level"]
        parts.append(
            f'<!-- wp:heading {{"level":{level}}} -->\n'
            f'<h{level}>{_html.escape(section["heading"])}</h{level}>\n'
            f"<!-- /wp:heading -->"
        )
        for block in re.split(r"\n\s*\n", section["markdown"] or ""):
            block = block.strip()
            if block:
                inner = _md_inline_to_html(block.replace("\n", " "))
                parts.append(f"<!-- wp:paragraph -->\n<p>{inner}</p>\n<!-- /wp:paragraph -->")
    return "\n\n".join(parts)


def _seo_meta(title: str, keyword: str, seo: dict | None) -> dict:
    """RankMath + Yoast SEO field mapping (PRD §15)."""
    seo = seo or {}
    focus = seo.get("focus_keyword") or keyword
    desc = seo.get("meta_description") or ""
    seo_title = seo.get("title") or title
    return {
        "rank_math_focus_keyword": focus,
        "rank_math_title": seo_title,
        "rank_math_description": desc,
        "_yoast_wpseo_focuskw": focus,
        "_yoast_wpseo_title": seo_title,
        "_yoast_wpseo_metadesc": desc,
    }


# --------------------------------------------------------------------------- #
# FR-10.2 — DOCX (lazy dependency, degrade clearly)
# --------------------------------------------------------------------------- #
def to_docx(draft: dict) -> bytes:
    """Render the draft as a ``.docx`` document and return its bytes (FR-10.2).

    Uses ``python-docx`` if installed (imported lazily); raises a clear
    :class:`NotImplementedError` when absent so the module still imports.
    """
    try:
        from docx import Document  # type: ignore[import-not-found]
    except ImportError as exc:  # pragma: no cover - exercised only without the dep
        raise NotImplementedError(
            "DOCX export requires the 'python-docx' package. Install it "
            "(pip install python-docx) to enable .docx export."
        ) from exc

    import io

    sections = _sections(draft)
    title = _title(draft, None)

    document = Document()
    document.add_heading(title, level=0)
    for section in sections:
        document.add_heading(section["heading"], level=section["level"])
        for block in re.split(r"\n\s*\n", section["markdown"] or ""):
            block = block.strip()
            if block:
                document.add_paragraph(block.replace("\n", " "))

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# --------------------------------------------------------------------------- #
# PRD §15 — WordPress publish (real REST; dry-run without credentials)
# --------------------------------------------------------------------------- #
def wordpress_publish(
    draft: dict,
    project: dict,
    creds: dict | None = None,
    *,
    schedule: str | None = None,
    seo: dict | None = None,
    status: str | None = None,
    featured_image: dict | None = None,
) -> dict:
    """Publish the draft to WordPress via the REST API (PRD §15, FR-10.2).

    Posts Gutenberg block content with RankMath/Yoast SEO meta and optional
    scheduling to ``{site}/wp-json/wp/v2/posts`` using an application-password
    (HTTP Basic auth). When credentials are missing it performs **no network
    call** and returns the would-publish payload (dry run), so the pipeline runs
    without a live WordPress instance.

    Args:
        draft / project: the draft payload + project brief.
        creds: ``{site_url, username, app_password}`` (an application password).
        schedule: ISO-8601 datetime to schedule the post (sets status ``future``).
        seo: ``{focus_keyword, meta_description, title}`` overrides.
        status: explicit WP post status (``publish`` / ``draft`` / ``future``).
        featured_image: ``{bytes, mime, filename, alt}`` — uploaded to the WP media
            library and set as the post's featured image (best-effort).

    Returns:
        ``{status: published|scheduled|would_publish|error, ...}``.
    """
    creds = creds or {}
    site_url = (creds.get("site_url") or creds.get("url") or "").rstrip("/")
    username = creds.get("username") or creds.get("user")
    app_password = creds.get("app_password") or creds.get("password")

    title = _title(draft, project)
    keyword = (project or {}).get("keyword") or ""
    content = to_gutenberg(draft)
    post_status = status or ("future" if schedule else "draft")

    payload: dict = {
        "title": title,
        "content": content,
        "status": post_status,
        "meta": _seo_meta(title, keyword, seo),
    }
    if schedule:
        payload["date"] = schedule  # ISO-8601; with status=future this schedules it
    endpoint = f"{site_url}/wp-json/wp/v2/posts" if site_url else "/wp-json/wp/v2/posts"

    # Dry run when credentials are absent — no network call (PRD §15).
    if not (site_url and username and app_password):
        return {
            "status": "would_publish",
            "note": "No WordPress credentials configured — dry run. Set WORDPRESS_* in .env to publish live.",
            "endpoint": endpoint,
            "payload": payload,
            "has_featured_image": bool(featured_image and featured_image.get("bytes")),
        }

    import base64

    import httpx

    from app.net import guard_ssrf

    # Re-validate at publish time (defense in depth vs DNS-rebinding), then never
    # follow redirects — a public host must not bounce us onto an internal one.
    guard_ssrf(site_url)

    auth = base64.b64encode(f"{username}:{app_password}".encode()).decode()
    headers = {"Authorization": f"Basic {auth}"}

    # Best-effort: upload the featured image to the media library and attach it.
    media_id = None
    if featured_image and featured_image.get("bytes"):
        try:
            filename = featured_image.get("filename") or "featured.png"
            media = httpx.post(
                f"{site_url}/wp-json/wp/v2/media",
                content=featured_image["bytes"],
                headers={
                    **headers,
                    "Content-Type": featured_image.get("mime") or "image/png",
                    "Content-Disposition": f'attachment; filename="{filename}"',
                },
                timeout=60,
                follow_redirects=False,
            )
            media.raise_for_status()
            media_id = media.json().get("id")
            if media_id:
                payload["featured_media"] = media_id
                if featured_image.get("alt"):
                    httpx.post(
                        f"{site_url}/wp-json/wp/v2/media/{media_id}",
                        json={"alt_text": featured_image["alt"]},
                        headers=headers,
                        timeout=30,
                        follow_redirects=False,
                    )
        except Exception:  # noqa: BLE001 - a media failure must not block the post
            media_id = None

    try:
        resp = httpx.post(endpoint, json=payload, headers=headers, timeout=30, follow_redirects=False)
        resp.raise_for_status()
        data = resp.json()
        return {
            "status": "scheduled" if post_status == "future" else "published",
            "id": data.get("id"),
            "link": data.get("link"),
            "wp_status": data.get("status"),
            "featured_media": media_id,
            "endpoint": endpoint,
        }
    except Exception as exc:  # noqa: BLE001 - surface any WP/transport error
        return {"status": "error", "error": str(exc), "endpoint": endpoint, "payload": payload}


def _slugify(text: str) -> str:
    import re

    slug = re.sub(r"[^a-z0-9]+", "-", (text or "").lower()).strip("-")
    return slug[:80] or "post"


def webhook_publish(
    draft: dict, project: dict | None, config: dict | None, *, status: str | None = None
) -> dict:
    """POST the finished article to a user's custom endpoint (generic publisher).

    For non-WordPress stacks: sends a stable JSON envelope (title, slug, HTML,
    Markdown, meta) that any backend can consume, with an optional
    ``Authorization: Bearer <token>``. No endpoint configured => a **dry run** (no
    network), so the pipeline still runs. SSRF-guarded and non-redirecting like the
    WordPress path. A 2xx response is success; the returned URL/id (if the endpoint
    sends JSON with ``url``/``link``/``id``) is surfaced back.
    """
    config = config or {}
    endpoint = (config.get("endpoint_url") or "").strip().rstrip("/")
    token = config.get("auth_token")
    post_status = status or config.get("default_status") or "draft"

    title = _title(draft, project)
    keyword = (project or {}).get("keyword") or ""
    envelope = {
        "source": "contentos",
        "status": post_status,
        "title": title,
        "slug": _slugify(title),
        "keyword": keyword,
        "html": to_html(draft),
        "markdown": to_markdown(draft),
        "meta": {"title": title, "focus_keyword": keyword},
    }

    if not endpoint:
        return {
            "status": "would_publish",
            "note": "No custom endpoint configured — dry run. Add one in Settings.",
            "endpoint": endpoint,
            "payload": envelope,
        }

    import httpx

    from app.net import guard_ssrf

    guard_ssrf(endpoint)
    headers = {"Content-Type": "application/json", "User-Agent": "ContentOS/1.0"}
    if token:
        headers["Authorization"] = f"Bearer {token}"
    try:
        resp = httpx.post(
            endpoint, json=envelope, headers=headers, timeout=30, follow_redirects=False
        )
        resp.raise_for_status()
        link = rid = None
        try:
            data = resp.json()
            if isinstance(data, dict):
                link = data.get("url") or data.get("link") or data.get("permalink")
                rid = data.get("id")
        except Exception:  # noqa: BLE001 - a non-JSON 2xx body is still a success
            pass
        return {
            "status": "published",
            "http_status": resp.status_code,
            "link": link,
            "id": rid,
            "endpoint": endpoint,
        }
    except httpx.HTTPStatusError as exc:
        return {
            "status": "error",
            "error": f"endpoint returned {exc.response.status_code}",
            "endpoint": endpoint,
        }
    except Exception as exc:  # noqa: BLE001 - surface any transport error
        return {"status": "error", "error": str(exc), "endpoint": endpoint}


# --------------------------------------------------------------------------- #
# PRD §15 — Google Docs export ("formatted document for review")
# --------------------------------------------------------------------------- #
def google_doc_export(draft: dict, project: dict, token: str | None = None) -> dict:
    """Export the draft to a Google Doc for review (PRD §15).

    With a Google OAuth access token it creates a Doc via the Docs API and inserts
    the formatted text; without a token it returns the doc-ready content (dry run)
    so the flow works before credentials are wired.
    """
    title = _title(draft, project)
    lines: list[str] = [title, ""]
    for section in _sections(draft):
        lines.append(section["heading"])
        for block in re.split(r"\n\s*\n", section["markdown"] or ""):
            block = block.strip()
            if block:
                lines.append(block.replace("\n", " "))
        lines.append("")
    body_text = "\n".join(lines).strip() + "\n"

    if not token:
        return {
            "status": "would_create",
            "title": title,
            "note": "No Google token configured — returning doc-ready content. Set GOOGLE_OAUTH_TOKEN to create the Doc.",
            "content": body_text,
        }

    import httpx

    headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}
    try:
        created = httpx.post(
            "https://docs.googleapis.com/v1/documents", json={"title": title}, headers=headers, timeout=30
        )
        created.raise_for_status()
        doc_id = created.json().get("documentId")
        httpx.post(
            f"https://docs.googleapis.com/v1/documents/{doc_id}:batchUpdate",
            json={"requests": [{"insertText": {"location": {"index": 1}, "text": body_text}}]},
            headers=headers,
            timeout=30,
        ).raise_for_status()
        return {
            "status": "created",
            "documentId": doc_id,
            "url": f"https://docs.google.com/document/d/{doc_id}/edit",
        }
    except Exception as exc:  # noqa: BLE001
        return {"status": "error", "error": str(exc)}
